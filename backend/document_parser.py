"""
document_parser.py
Parse multiple document types into plain text for LightRAG ingestion.
"""

import os
from typing import Optional


def parse_markdown(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_pdf(path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def parse_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            row_text = "\t".join(cells).strip()
            if row_text:
                lines.append(row_text)
    wb.close()
    return "\n".join(lines)


def parse_image(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        err_name = type(e).__name__
        if "TesseractNotFound" in err_name or "tesseract" in str(e).lower():
            return ""  # Tesseract not installed — silently skip
        raise


PARSERS = {
    ".md": parse_markdown,
    ".txt": parse_markdown,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".png": parse_image,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".tiff": parse_image,
    ".tif": parse_image,
    ".bmp": parse_image,
    ".webp": parse_image,
}


def parse_document(path: str) -> Optional[str]:
    """Return extracted text or None if the extension is unsupported."""
    ext = os.path.splitext(path)[1].lower()
    parser = PARSERS.get(ext)
    if parser is None:
        return None
    try:
        text = parser(path)
        return text if text and text.strip() else None
    except Exception as e:
        raise RuntimeError(f"Failed to parse {path}: {e}") from e
